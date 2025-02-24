import json
import os.path
from dataclasses import dataclass
from re import Match
from typing import Iterator
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx2pdf import convert
from flet.utils import deprecated

from core.global_handlers import PLACEHOLDERS_FOLDER, LOGGER
from utils.json_import_export import import_json, save_json
from utils.path_utils import *


@dataclass
class Placeholder:
    text: str
    name: str
    number: int
    tooltip: str = None
    new_text: str = None

    def to_dict(self):
        return {
            'text':     self.text,
            'name':     self.name,
            'number':   self.number,
            'tooltip':  self.tooltip,
            'new_text': self.new_text
        }


def _remove_ph_escaped_text(input_string: str) -> str:
    """
    Removes all text enclosed by || || from the input string.

    :param input_string (str): The string to process.
    :return: A string with text inside || || removed.
    """
    # Use regex to match ||text|| and replace it with an empty string
    return re.sub(r'\|.*?\|', '', input_string).strip()


class DocManager:
    def __init__(self, template_name: Union[str | Path]):
        self.save_docx_path = ""
        self.pm = PathManager(template_name)
        self.template_path = self.pm.resolve_new_path
        year = datetime.now().strftime("%Y")
        month = datetime.now().strftime("%m-%b")
        self.output_dir = PathManager(f"/docs/Applications/{year}/{month}",
                                      PathFlag.CREATE_FOLDER | PathFlag.FROM_PROJECT_ROOT)
        self.doc = Document(str(self.template_path)) if self.template_path else None
        self.placeholders: dict = {}

        self._json_dir: str = create_folder_if_dne(get_project_root() + f"/docs/Applications/{year}/{month}")

    @property
    def callable(self) -> bool:
        return bool(self.doc)

    def _json_path(self, create_if_not_exist=False):
        """
        Generate the path to the JSON file associated with the document.
        """
        if self.save_docx_path == "":
            json_path = normalize_path(f"{get_project_root()}/placeholders/placeholders.json")
        else:
            json_name = os.path.basename(self.save_docx_path).replace('.docx', '.json')
            json_path = normalize_path(os.path.join(self._json_dir, json_name))

        if create_if_not_exist and not os.path.exists(json_path):
            try:
                with open(json_path, "w") as f:
                    json.dump({}, f)
            except Exception as e:
                LOGGER.log(e)
                print(f"Failed to create JSON file at {json_path}: {e}")

        return json_path

    def _map_placeholders(self, paragraph):
        """
        @APPROVED
        Helper method to process a paragraph or table cell content.
        """
        runs = paragraph.text
        placeholders = self._find_placeholders(runs)
        if placeholders:
            for p in placeholders:
                placeholder_text = p.group(0)  # Extract the matched string
                if placeholder_text.startswith('[['):  # Check for '[[' at the start
                    self.placeholders[placeholder_text] = placeholder_text[2:-2]
                else:  # e.g., if p.startswith('{{') and p.endswith('}}')
                    self.placeholders[placeholder_text] = ''

    def get_placeholders(self, force_refresh=False, update_json=False):
        if not self.callable:
            return
        if len(self.placeholders) <= 0 or force_refresh:
            self._update_placeholders()
        return self.placeholders

    def _fill_empty_placeholders(self):
        old_placeholders = self._import_json()
        old_keys = old_placeholders.keys()
        if not old_keys or len(old_keys) == 0:
            return self.placeholders
        for key, value in self.placeholders.items():
            if value == '' or value is None:
                if key in old_keys:
                    self.placeholders[key] = old_placeholders[key]

    def _update_placeholders(self, fill_empty_placeholders: bool = False):
        """
        Gets the placeholders from the DocX file and stores them in memory.
        Optional: refresh the JSON file.
        """
        contents = self.doc.iter_inner_content()
        self.placeholders = {}
        # Header + Footer
        if self.has_header(self.doc.sections[0]):
            if self.doc.sections[0].header:
                header_text = self.doc.sections[0].header.paragraphs
                for paragraph in header_text:
                    self._map_placeholders(paragraph)
        if self.has_footer(self.doc.sections[0]):
            if self.doc.sections[0].footer:
                footer_text = self.doc.sections[0].footer.paragraphs
                for paragraph in footer_text:
                    self._map_placeholders(paragraph)

        # Paragraphs + Tables
        for item in contents:
            if isinstance(item, Paragraph):
                self._map_placeholders(item)
            elif isinstance(item, Table):
                for row in item.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._map_placeholders(paragraph)

        if fill_empty_placeholders:
            self._fill_empty_placeholders()

    def apply_replacements(self, new_values, save_placeholders=True):
        """
        @APPROVED
        Replace placeholders with values from a dictionary."""
        if not self.callable:
            return
        for paragraph in self.doc.paragraphs:
            self._replace_in_paragraph(paragraph, new_values)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, new_values)

        # Header + Footer

        if self.has_header(self.doc.sections[0]):
            header_text = self.doc.sections[0].header.paragraphs
            for paragraph in header_text:
                self._replace_in_paragraph(paragraph, new_values)
        if self.has_footer(self.doc.sections[0]):
            footer_text = self.doc.sections[0].footer.paragraphs
            for paragraph in footer_text:
                self._replace_in_paragraph(paragraph, new_values)

        if save_placeholders:
            self.save_placeholders_to_json()

    def _replace_in_paragraph(self, paragraph, new_values):
        """Process a single paragraph to replace placeholders across runs."""
        runs = paragraph.runs
        if not runs:
            return

        # Build a list of run texts and their cumulative positions
        run_texts = [run.text for run in runs]
        full_text = ''.join(run_texts)

        # Find all placeholders
        matches = list(self._find_placeholders(full_text))

        # Iterate over matches in reverse order to prevent index shifting
        for match in reversed(matches):
            placeholder = match.group(0)
            if placeholder not in new_values:
                continue
            replacement = new_values[placeholder]

            start, end = match.start(), match.end()

            # Find the runs that cover this placeholder
            run_start_idx, run_end_idx, offset_start, offset_end = self._find_run_indices(run_texts, start, end)
            if run_start_idx is None:
                continue  # Could not find runs covering this placeholder

            # Replace text in runs
            self._replace_text_in_runs(runs, run_start_idx, run_end_idx, offset_start, offset_end,
                                       replacement)

    @staticmethod
    def _find_run_indices(run_texts, start, end):
        """Find the run indices and offsets that cover the placeholder."""
        current_pos = 0
        run_start_idx = run_end_idx = None
        offset_start = offset_end = 0
        for i, text in enumerate(run_texts):
            run_length = len(text)
            run_end_pos = current_pos + run_length
            if run_start_idx is None and start < run_end_pos:
                run_start_idx = i
                offset_start = start - current_pos
            if run_end_idx is None and end <= run_end_pos:
                run_end_idx = i
                offset_end = end - current_pos
                break
            current_pos += run_length
        if run_start_idx is not None and run_end_idx is not None:
            return run_start_idx, run_end_idx, offset_start, offset_end
        else:
            return None, None, None, None

    @staticmethod
    def _replace_text_in_runs(runs, run_start_idx, run_end_idx, offset_start, offset_end, replacement: str):

        """Replace the placeholder text in the specified runs with the replacement."""
        # Collect runs involved in the placeholder
        affected_runs = runs[run_start_idx:run_end_idx + 1]

        # Handle the case where the placeholder spans multiple runs
        if run_start_idx < run_end_idx:
            # Modify the first run
            first_run = affected_runs[0]
            before_placeholder = first_run.text[:offset_start]
            first_run.text = before_placeholder + _remove_ph_escaped_text(replacement)

            # Modify the last run
            last_run = affected_runs[-1]
            after_placeholder = last_run.text[offset_end:]
            last_run.text = after_placeholder

            # Clear intermediate runs
            for run in affected_runs[1:-1]:
                if not run.text == '':
                    run._element.getparent().remove(run._element)

        else:  # Single run case
            single_run = affected_runs[0]
            before_placeholder = single_run.text[:offset_start]
            after_placeholder = single_run.text[offset_end:]
            single_run.text = before_placeholder + _remove_ph_escaped_text(replacement) + after_placeholder

    def replace_hyperlink(self, old_link, old_text, new_link, new_text):
        """
        @APPROVED
        Replace all hyperlinks with target `old_link` and display text `old_text`
        to point to `new_link` and update the displayed text to `new_text`.

        Checks headers, footers, tables, and paragraphs, and handles hyperlinks spanning multiple runs.

        :param old_link: The hyperlink target to search for.
        :param old_text: The displayed text of the hyperlink to search for.
        :param new_link: The new hyperlink target to replace with.
        :param new_text: The new text to display for the hyperlink.
        """
        if not self.callable:
            return
        namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        rel_namespace = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

        def process_xml_hyperlinks(container):
            """ @Approved """

            for paragraph in container.paragraphs:
                process_xml_paragraph(paragraph)

            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_xml_hyperlinks(cell)  # Recursive call for nested structures

        def process_xml_paragraph(paragraph):
            """ @Approved """
            p = paragraph._p  # CT_P (paragraph XML element)
            hyperlinks = p.findall('.//w:hyperlink', namespace)

            for hyperlink in hyperlinks:
                rId = hyperlink.get(f'{{{rel_namespace}}}id')
                if not rId:
                    continue

                rels = self.doc.part.rels
                if rId in rels and rels[rId]._target == old_link:
                    # Collect all run text under this hyperlink
                    text_runs = []
                    for run in hyperlink.findall('.//w:r', namespace):
                        run_text = run.find('.//w:t', namespace)
                        if run_text is not None:
                            text_runs.append(run_text)

                    # Combine text from runs and check if it matches old_text
                    full_text = ''.join(rt.text for rt in text_runs)
                    if full_text == old_text:
                        rels[rId]._target = new_link
                        # Update the text while preserving formatting
                        for i, run_text in enumerate(text_runs):
                            if i == 0:
                                run_text.text = new_text
                            else:
                                run_text.text = ''

        process_xml_hyperlinks(self.doc)
        for section in self.doc.sections:
            if section.header:
                process_xml_hyperlinks(section.header)
            if section.footer:
                process_xml_hyperlinks(section.footer)

    def save_docx(self, output_name):
        if not self.callable:
            return

        self.save_docx_path = str(self.output_dir.resolve_new_path.with_name(output_name))
        self.doc.save(self.save_docx_path)
        LOGGER.log('Saved to ' + self.save_docx_path)
        return self.save_docx_path

    def save_pdf(self, output_name: str):
        """
        save the .docx as a PDF.
        """
        if not self.callable:
            return
        rename_file_by_creation(PathManager.resolve_path('/PDF Output/' + output_name, PathFlag.R | PathFlag.C))
        out = normalize_path(get_project_root() + '/PDF Output/' + output_name)
        LOGGER.log('Saved to ' + out)
        convert(self.save_docx_path, out)
        return out

    def save_placeholders_to_json(self):
        """
        Save the extracted placeholders to a JSON file.
        """
        path = PLACEHOLDERS_FOLDER.resolve_new_path
        path = path.with_name('_ph_' + self.pm.new_path_name)
        path = path.with_suffix('.json')
        PathManager.resolve_path(path,
                                 PathFlag.CREATE_FOLDER | PathFlag.INCREMENT_IF_EXISTS | PathFlag.CASCADE_BY_YEAR | PathFlag.CASCADE_BY_MONTH)
        if not self.callable:
            return
        return save_json(self.placeholders, path, False, True)

    def _import_json(self):
        return import_json(self._json_path(False))

    @staticmethod
    def _find_placeholders(paragraph: str) -> Iterator[Match[str]]:
        """ @Approved
        Return values includes both square and curly braces.
        Returned as list; in order of appearance.
        Extract placeholders and defaults from a paragraph's text.
        Placeholders: surrounded by `{{...}}`.
        Defaults: surrounded by `[[...]]`.
        Nested not allowed.
        Auto-escape trailing extra braces/brackets.
        """
        pattern = r'{{([^}]*)}}+|\[\[([^\]]*)\]\]+'
        return re.finditer(pattern, paragraph)

    @staticmethod
    def _create_documents_folder(employer_name: str) -> str:
        """
        FIXME: NOT USED - SHOULD BE USED
        Creates a folder structure for organizing job application documents based on
        the current year and month (formatted as number-month). Includes a specific
        folder for the given employer. Checks if the folder exists.

        :param employer_name: The name of the employer for the folder.
        :return: The normalized path to the created folder.
        """
        now = datetime.today()
        year = now.year
        month = f"{now.month:02d}-{now.strftime('%B')}"  # Format: "01-January"
        base_folder = create_folder_if_dne('../docs_private')
        employer_folder = normalize_path(os.path.join(base_folder, str(year), month, employer_name))
        if not os.path.exists(employer_folder):
            os.makedirs(employer_folder)
        return employer_folder

    @staticmethod
    def has_header(section):
        """
        Determines if a header exists in the section without triggering its initialization.
        Checks the raw XML for a header relationship.
        """
        print('called has_header')
        return section.header == True

    @staticmethod
    def has_footer(section):
        """
        Determines if a footer exists in the section without triggering its initialization.
        Checks the raw XML for a footer relationship.
        """
        print('called has_footer')
        return section.footer == True

    @deprecated('not needed; use _import_json instead', version='2025-01-10', delete_version='Not sure yet')
    def _load_json(self, force_update: bool = False):
        """
        Load placeholders from JSON and merge with existing placeholders.
        Retain the old value if and only if the new value is empty.
        Overwrites old placeholders if they overlap \n
        **<!> Use with caution! <!>**
        """

        if not self.placeholders or force_update:
            self._import_json()
        else:
            old_placeholders = self.placeholders.copy()
            self._import_json()
            # old values are used if new ones are missing or empty strings.
            for key, old_value in old_placeholders.items():
                new_value = self.placeholders.get(key, "")
                if new_value == "":
                    self.placeholders[key] = old_value
        return self.placeholders


l1 = ["Python, Flet, Flask",
      "Algorithm analysis and design",
      "Object-oriented programming (OOP)",
      "Java, Kotlin",
      "Database design and optimization",
      "Automated testing (Pytest)",
      "C, C++",
      "Version Control, Git, BitBucket",
      "Agile & Scrum methodologies",
      "HTML, CSS, JavaScript",
      "Cross-functional collaboration",
      "Troubleshooting and debugging",
      "SQL, SQLite, MySQL",
      "Data structures and Graph Theory",
      "Requirements gathering"]
l2 = ["Python, Flet, Flask",
      "Algorithm analysis and design",
      "Object-oriented programming (OOP)",
      "Java, Kotlin",
      "Database design and optimization",
      "Automated testing (Pytest)",
      "C, C++",
      "Version Control, Git, BitBucket",
      "Agile & Scrum methodologies",
      "HTML, CSS, JavaScript",
      "Cross-functional collaboration",
      "Troubleshooting and debugging",
      "SQL, SQLite, MySQL",
      "Data structures and Graph Theory",
      "Requirements gathering"]

print(l1 == l2)


#
# #
# # Example Usage
# if __name__ == "__main__":
#     docx_path = normalize_path(get_project_root() + '/docs/templates/Jonathan Alexander - Resume.docx')
#     doc_man = DocManager(docx_path)
#     replacements = doc_man.get_placeholders()
#     # print(replacements)
#     for key, val in replacements.items():
#         replacements[key] = key.replace('{', '!').replace('}', '!').replace('[', '!').replace(']', '!')
#     # print(replacements)
#     doc_man.apply_replacements(replacements)
#     doc_man.save_docx("output.docx")
